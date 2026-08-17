"""Render a markdown report into the official submission template.

Loads the sprint template so page setup and paragraph styles are inherited,
clears its guidance text, and rebuilds the body from markdown.

    .venv\\Scripts\\python.exe scripts\\11_build_docx.py report/report_4page.md
    .venv\\Scripts\\python.exe scripts\\11_build_docx.py report/report2.md

Writes alongside the source as .docx. Open in Word or Google Docs and export to
PDF for submission, the sprint requires a PDF, and exporting from the template
keeps its fonts and margins.

Supported markdown: headings, paragraphs, bullet and numbered lists, tables,
blockquotes, horizontal rules, images via ![caption](path), and inline
**bold** / *italic* / `code`.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "template" / "Copy of Digital Minds Research Sprint submission template.docx"

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")

# Anything that begins a new block rather than continuing the current one.
BLOCK_START = re.compile(r"^(?:#{1,6}\s|>\s|[-*+]\s|\d+\.\s|\||```|!\[|---|===)")


def gather(lines: list[str], i: int) -> tuple[str, int]:
    """Join a hard-wrapped markdown block into one logical line.

    The sources are wrapped at about 100 columns so diffs stay readable. Without
    this, every wrapped line became its own Word paragraph: the 4-page report
    came out at thirteen pages of ragged text with orphan words stranded on their
    own lines. Word does the wrapping, so the source wrapping has to be undone.
    """
    parts = [lines[i].strip()]
    j = i + 1
    while j < len(lines):
        nxt = lines[j].strip()
        if not nxt or BLOCK_START.match(nxt):
            break
        parts.append(nxt)
        j += 1
    return " ".join(parts), j


def clear_body(doc: Document) -> None:
    """Strip the template's guidance content, keeping section setup and styles."""
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def add_runs(para, text: str) -> None:
    """Add text to a paragraph, honouring inline bold/italic/code."""
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            para.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            para.add_run(chunk[1:-1]).italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = para.add_run(chunk[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        else:
            para.add_run(chunk)


def has_style(doc: Document, name: str) -> bool:
    try:
        doc.styles[name]
        return True
    except KeyError:
        return False


def set_borders(table) -> None:
    """Apply single-line borders directly.

    The sprint template ships no 'Table Grid' style, so borders have to be set
    on the table element rather than via a named style.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "999999")
        borders.append(el)
    tbl_pr.append(borders)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    if has_style(doc, "Table Grid"):
        t.style = "Table Grid"
    else:
        set_borders(t)
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j >= len(t.columns):
                continue
            para = t.cell(i, j).paragraphs[0]
            add_runs(para, cell)
            for run in para.runs:
                run.font.size = Pt(8)
                if i == 0:
                    run.bold = True


def build(md_path: Path) -> Path:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document(str(TEMPLATE))
    clear_body(doc)

    i, in_code, table_buf = 0, False, []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            add_table(doc, table_buf)
            doc.add_paragraph()
            table_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = "Consolas"
            r.font.size = Pt(8)
            i += 1
            continue

        # Tables
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if not all(set(c) <= set("-: ") for c in cells):  # skip separator row
                table_buf.append(cells)
            i += 1
            continue
        flush_table()

        if not stripped:
            i += 1
            continue

        # Drop the draft-status blockquote and horizontal rules
        if stripped.startswith("---") or stripped.startswith("==="):
            i += 1
            continue

        # Images
        m = re.match(r"!\[(.*?)\]\((.+?)\)", stripped)
        if m:
            caption, rel = m.group(1), m.group(2)
            img = (ROOT / rel).resolve()
            if img.exists():
                doc.add_picture(str(img), width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cap.add_run(caption)
                r.italic = True
                r.font.size = Pt(8)
            else:
                print(f"  WARNING: image not found: {img}")
            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[2:].strip())
            r.bold = True
            r.font.size = Pt(16)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_paragraph(stripped[3:].strip(), style="Heading 2")
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:].strip(), style="Heading 3")
            i += 1
            continue

        # Blockquote, used for draft notes; render small and grey
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[2:])
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # Lists
        # Lists. The template defines no list styles, so indent manually and
        # write the marker into the text rather than relying on numbering.
        # List items and paragraphs all wrap across source lines, so each is
        # gathered before being written.
        if re.match(r"^(\d+)\.\s+", stripped):
            text, i = gather(lines, i)
            m = re.match(r"^(\d+)\.\s+(.*)", text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, f"{m.group(1)}. {m.group(2)}")
            continue
        if stripped.startswith("- "):
            text, i = gather(lines, i)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, f"• {text[2:]}")
            continue

        text, i = gather(lines, i)
        p = doc.add_paragraph()
        # Without this, consecutive paragraphs butt against each other and the
        # Results section reads as one unbroken wall of text.
        p.paragraph_format.space_after = Pt(6)
        add_runs(p, text)

    flush_table()

    out = md_path.with_suffix(".docx")
    doc.save(str(out))
    return out


def main() -> None:
    targets = [Path(a) for a in sys.argv[1:]] or [
        ROOT / "report" / "report_4page.md",
        ROOT / "report" / "report2.md",
    ]
    for t in targets:
        t = t if t.is_absolute() else ROOT / t
        if not t.exists():
            print(f"  missing: {t}")
            continue
        out = build(t)
        print(f"  {t.name} -> {out.relative_to(ROOT)}  ({out.stat().st_size/1024:.0f} KB)")


if __name__ == "__main__":
    main()
